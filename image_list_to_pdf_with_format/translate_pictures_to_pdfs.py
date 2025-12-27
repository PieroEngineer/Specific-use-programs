import os
from pathlib import Path
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from tempfile import TemporaryDirectory

def get_pdf_files(folder_path):
    """Returns a list of full paths to .pdf files in the folder"""
    return [str(f) for f in Path(folder_path).glob("*.pdf")]

def convert_first_page_to_image(pdf_path, output_folder):
    """Converts first page of a PDF to an image and returns the image path"""
    images = convert_from_path(pdf_path, first_page=1, last_page=1, fmt="png", output_folder=output_folder, poppler_path = r'poppler-24.08.0\Library\bin')
    return images[0].filename if images else None

def insert_images_into_word(image_paths, output_doc_path):
    """Inserts a list of images into a Word document, one per page"""
    doc = Document()
    section = doc.sections[0] # Get the first section

    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

    header = section.header
    footer = section.footer

    # For header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_path = r'standart\header.png'
    header_paragraph.add_run().add_picture(header_path, width=Inches(9.0))

    # For footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_path = r'standart\footer.png'
    footer_paragraph.add_run().add_picture(footer_path, width=Inches(9.0))

    for i, image_path in enumerate(image_paths):
        p = doc.add_paragraph()
        p.alignment = 1  ## 0 = left, 1 = center, 2 = right
        p.add_run().add_picture(image_path, width=Inches(5))

        if i < len(image_paths) - 1:
            doc.add_page_break()
    doc.save(output_doc_path)
    return output_doc_path

def process_folder_to_word(folder_path, output_doc_name="document_file.docx"):
    """Main function: extracts first page images from PDFs and creates a Word doc"""
    pdf_files = get_pdf_files(folder_path)
    if not pdf_files:
        raise FileNotFoundError("No PDF files found in the provided folder.")
    
    output_folder_path = r'Resultado'
    output_doc_path = os.path.join(output_folder_path, output_doc_name)

    with TemporaryDirectory() as temp_dir:
        image_paths = []
        for pdf in pdf_files:
            img_path = convert_first_page_to_image(pdf, temp_dir)
            if img_path:
                image_paths.append(img_path)
        insert_images_into_word(image_paths, output_doc_path)

    return output_doc_path

# 🔁 Example usage
if __name__ == "__main__":
    #folder = input("📁 Enter folder path with PDFs: ").strip('"')
    folder = r'Entrada'
    try:
        result = process_folder_to_word(folder)
        print(f"✅ Word document created at:\n{result}")
    except Exception as e:
        print(f"❌ Error: {e}")
